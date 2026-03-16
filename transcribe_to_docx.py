import json
import os
import subprocess
import sys
import wave
from pathlib import Path

import imageio_ffmpeg
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from vosk import KaldiRecognizer, Model, SetLogLevel


ROOT = Path(__file__).resolve().parent
VIDEO_EXTENSIONS = {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".webm"}
MODEL_DIR = ROOT / "vosk-model-small-ru-0.22"


def find_video() -> Path:
    videos = sorted(
        path for path in ROOT.iterdir() if path.is_file() and path.suffix.lower() in VIDEO_EXTENSIONS
    )
    if not videos:
        raise FileNotFoundError("Video file not found in the working directory.")
    return videos[0]


def extract_audio(video_path: Path, wav_path: Path) -> None:
    ffmpeg = imageio_ffmpeg.get_ffmpeg_exe()
    cmd = [
        ffmpeg,
        "-y",
        "-i",
        str(video_path),
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        "-sample_fmt",
        "s16",
        str(wav_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)


def read_results(recognizer: KaldiRecognizer, audio_file: wave.Wave_read):
    results = []
    while True:
        data = audio_file.readframes(4000)
        if len(data) == 0:
            break
        if recognizer.AcceptWaveform(data):
            payload = json.loads(recognizer.Result())
            if payload.get("result"):
                results.extend(payload["result"])
    final_payload = json.loads(recognizer.FinalResult())
    if final_payload.get("result"):
        results.extend(final_payload["result"])
    return results


def group_words(words):
    segments = []
    current = []
    max_words = 22
    pause_threshold = 0.9

    for word in words:
        if current:
            prev = current[-1]
            pause = word["start"] - prev["end"]
            if pause > pause_threshold or len(current) >= max_words:
                segments.append(current)
                current = []
        current.append(word)

    if current:
        segments.append(current)

    return segments


def format_ts(seconds: float) -> str:
    total = int(seconds)
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def sentence_case(text: str) -> str:
    text = " ".join(text.split())
    if not text:
        return text
    return text[0].upper() + text[1:]


def build_transcript(words):
    segments = group_words(words)
    transcript_segments = []
    for segment in segments:
        start = format_ts(segment[0]["start"])
        end = format_ts(segment[-1]["end"])
        text = sentence_case(" ".join(item["word"] for item in segment))
        transcript_segments.append((start, end, text))
    return transcript_segments


def save_txt(output_path: Path, video_name: str, segments) -> None:
    lines = [
        "Подробная транскрибация видео",
        f"Файл: {video_name}",
        "",
        "Примечание: документ создан автоматически офлайн; возможны неточности в пунктуации и именах.",
        "",
    ]
    for start, end, text in segments:
        lines.append(f"[{start} - {end}] {text}")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def save_docx(output_path: Path, video_name: str, duration_label: str, segments) -> None:
    document = Document()

    base_style = document.styles["Normal"]
    base_style.font.name = "Times New Roman"
    base_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Подробная транскрибация видео")
    title_run.bold = True
    title_run.font.size = Pt(14)

    meta = document.add_paragraph()
    meta.add_run("Файл: ").bold = True
    meta.add_run(video_name)
    meta.add_run("\nДлительность: ").bold = True
    meta.add_run(duration_label)
    meta.add_run("\nПримечание: ").bold = True
    meta.add_run("автоматическая офлайн-транскрибация; возможны неточности в пунктуации и именах.")

    document.add_paragraph("")

    for start, end, text in segments:
        paragraph = document.add_paragraph()
        stamp = paragraph.add_run(f"[{start} - {end}] ")
        stamp.bold = True
        paragraph.add_run(text)

    document.save(str(output_path))


def main():
    if not MODEL_DIR.exists():
        raise FileNotFoundError(f"Model not found: {MODEL_DIR}")

    video_path = find_video()
    wav_path = ROOT / f"{video_path.stem}.wav"
    txt_path = ROOT / f"{video_path.stem}.transcript.txt"
    docx_path = ROOT / f"{video_path.stem}.transcript.docx"

    print(f"Extracting audio from {video_path.name}...")
    extract_audio(video_path, wav_path)

    SetLogLevel(-1)
    model = Model(str(MODEL_DIR))

    with wave.open(str(wav_path), "rb") as audio_file:
        recognizer = KaldiRecognizer(model, audio_file.getframerate())
        recognizer.SetWords(True)
        words = read_results(recognizer, audio_file)

    if not words:
        raise RuntimeError("No speech recognized in the audio track.")

    segments = build_transcript(words)
    duration_label = segments[-1][1]

    save_txt(txt_path, video_path.name, segments)
    save_docx(docx_path, video_path.name, duration_label, segments)

    print(f"TXT saved to: {txt_path}")
    print(f"DOCX saved to: {docx_path}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        sys.exit(1)
